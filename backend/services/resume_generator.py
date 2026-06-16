"""
Resume Generator — Builds Premium, ATS-friendly DOCX resumes.
Optimized for enterprise standards and professional aesthetics.
"""

import io
import os
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


OUTPUT_DIR = Path(__file__).parent.parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

TEMPLATE_PATH = Path(__file__).parent.parent / "data" / "templates.json"

# Maps template section_order keys → structured_data["sections"] keys
SECTION_MAP = {
    "work_experience": "work_experience",
    "internships": "internships",
    "technical_skills": "technical_skills",
    "projects": "projects",
    "education": "education",
    "certifications": "certifications",
    "awards": "awards",
    "interests": "interests",
    "summary": "__summary__",       # special: from candidate.summary
    "core_competencies": "__summary__",  # same source, avoid duplicate
}

# Maps SECTION_MAP result → final_sections (optimizer) internal key
FINAL_KEY_MAP = {
    "work_experience": "experience",
    "internships": "internships",
    "technical_skills": "skills",
    "projects": "projects",
    "education": "education",
    "certifications": "certifications",
    "awards": "achievements",
    "interests": "interests",
    "__summary__": "summary",
}


def _get_template_config(template_type: str) -> dict:
    try:
        with open(TEMPLATE_PATH, "r") as f:
            data = json.load(f)
            for t in data.get("templates", []):
                if t["template_type"] == template_type:
                    return t
    except Exception:
        pass
    return {
        "rendering": {"font_family": "Arial", "font_size_body": 11, "font_size_heading": 12},
        "preferred_section_order": ["summary", "work_experience", "technical_skills", "projects", "education", "certifications"]
    }


def _add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E3A59')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_heading(doc, text: str, rendering: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = rendering.get("font_family", "Arial")
    run.font.size = Pt(rendering.get("font_size_heading", 12))
    run.font.color.rgb = RGBColor(0x2E, 0x3A, 0x59)
    _add_horizontal_line(p)


def _add_bullet(doc, text: str, rendering: dict):
    text = text.strip().lstrip("•-–*▪►→").strip()
    if not text:
        return
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"• {text}")
    run.font.name = rendering.get("font_family", "Arial")
    run.font.size = Pt(rendering.get("font_size_body", 11))


def _render_job_entry(doc, job: dict, rendering: dict):
    """Render a single work experience or internship entry."""
    font_family = rendering.get("font_family", "Arial")
    font_size_body = rendering.get("font_size_body", 11)

    job_title = job.get("job_title", "").strip()
    employer = job.get("employer", "").strip()
    start = job.get("start_date", "").strip()
    end = job.get("end_date", "").strip()
    location = job.get("location", "").strip()
    summary = job.get("summary", "").strip()
    achievements = job.get("achievements", [])

    if not job_title:
        return

    # Title + Date row
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(1)
    r1 = p1.add_run(job_title)
    r1.bold = True
    r1.font.size = Pt(font_size_body + 0.5)
    r1.font.name = font_family

    date_str = ""
    if start and end:
        date_str = f"{start} – {end}"
    elif end:
        date_str = end
    elif start:
        date_str = f"{start} – Present"

    if date_str:
        dr = p1.add_run(f"    {date_str}")
        dr.bold = True
        dr.font.size = Pt(font_size_body)
        dr.font.name = font_family

    # Company + Location row
    if employer:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        cr = p2.add_run(employer)
        cr.italic = True
        cr.font.size = Pt(font_size_body)
        cr.font.name = font_family
        if location:
            lr = p2.add_run(f" | {location}")
            lr.italic = True
            lr.font.size = Pt(font_size_body)
            lr.font.name = font_family

    # Summary sentence
    if summary:
        _add_bullet(doc, summary, rendering)

    # Achievement bullets
    for ach in (achievements or []):
        if ach.strip():
            _add_bullet(doc, ach, rendering)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _render_education_entry(doc, edu: dict, rendering: dict):
    font_family = rendering.get("font_family", "Arial")
    font_size_body = rendering.get("font_size_body", 11)

    degree = edu.get("degree", "").strip()
    institution = edu.get("institution", "").strip()
    start = edu.get("start_date", "").strip()
    end = edu.get("end_date", "").strip()
    cgpa = edu.get("cgpa", "").strip()

    if not institution and not degree:
        return

    # Degree + Date
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(1)
    label = degree if degree else institution
    r1 = p1.add_run(label)
    r1.bold = True
    r1.font.size = Pt(font_size_body + 0.5)
    r1.font.name = font_family

    date_str = ""
    if start and end:
        date_str = f"{start} – {end}"
    elif end:
        date_str = end

    if date_str:
        dr = p1.add_run(f"    {date_str}")
        dr.bold = True
        dr.font.size = Pt(font_size_body)
        dr.font.name = font_family

    # Institution (if degree was shown above)
    if degree and institution:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        ir = p2.add_run(institution)
        ir.italic = True
        ir.font.size = Pt(font_size_body)
        ir.font.name = font_family
        if cgpa:
            p2.add_run(f" | CGPA: {cgpa}").italic = True
    elif cgpa:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.add_run(f"CGPA: {cgpa}").italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _render_project_entry(doc, proj: dict, rendering: dict):
    font_family = rendering.get("font_family", "Arial")
    font_size_body = rendering.get("font_size_body", 11)

    name = proj.get("project_name") or proj.get("title", "")
    description = proj.get("description", "")
    technologies = proj.get("technologies") or proj.get("tech_stack", [])
    # Handle 3-bullet format from suggest_projects
    bullets = proj.get("bullets", [])

    if not name:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(name.strip())
    r.bold = True
    r.font.size = Pt(font_size_body + 0.5)
    r.font.name = font_family

    if technologies:
        tr = p.add_run(f" | {', '.join(technologies)}")
        tr.italic = True
        tr.font.size = Pt(font_size_body)
        tr.font.name = font_family

    if bullets:
        for bullet in bullets:
            _add_bullet(doc, bullet, rendering)
    elif description:
        _add_bullet(doc, description, rendering)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx(
    final_sections: dict,
    committed_projects: list,
    candidate_name: str = "Candidate",
    template: str = "chronological_classic",
    structured_data: dict = None,
) -> bytes:
    """Build Premium ATS-friendly DOCX."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed.")

    config = _get_template_config(template)
    rendering = config.get("rendering", {})
    section_order = config.get("preferred_section_order", [])
    font_family = rendering.get("font_family", "Arial")
    font_size_body = rendering.get("font_size_body", 11)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── HEADER ─────────────────────────────────────────────────────────────
    sd = structured_data or {}
    header_data = sd.get("candidate", {})
    name = header_data.get("full_name", candidate_name).strip().upper()

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.name = font_family
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    name_para.paragraph_format.space_after = Pt(2)

    contact = []
    if header_data.get("email"):     contact.append(header_data["email"])
    if header_data.get("phone"):     contact.append(header_data["phone"])
    loc = header_data.get("location", {})
    if loc.get("city"):              contact.append(f"{loc['city']}, {loc.get('state', '')}".strip(", "))
    if header_data.get("linkedin_url"): contact.append(header_data["linkedin_url"])
    if header_data.get("github_url"):   contact.append(header_data["github_url"])

    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(" | ".join(contact))
        cr.font.size = Pt(9.5)
        cr.font.name = font_family
        cp.paragraph_format.space_after = Pt(10)

    # ── SECTIONS ───────────────────────────────────────────────────────────
    rendered_summary = False  # Prevent summary from appearing twice

    for json_key in section_order:
        struct_key = SECTION_MAP.get(json_key, json_key)
        final_key = FINAL_KEY_MAP.get(struct_key, struct_key)

        # Summary — render only once
        if struct_key == "__summary__":
            if rendered_summary:
                continue
            rendered_summary = True
            summary_text = (
                sd.get("candidate", {}).get("summary", "")
                or final_sections.get("summary", {}).get("content", "")
            ).strip()
            if not summary_text:
                continue
            _add_heading(doc, "Professional Summary", rendering)
            _add_bullet(doc, summary_text, rendering)
            continue

        # Get data from structured_data sections
        sections_data = sd.get("sections", {})
        struct_data = sections_data.get(struct_key)
        has_struct = struct_data is not None and struct_data != [] and struct_data != ""
        has_committed_projects = (struct_key == "projects" and committed_projects)
        final_content = final_sections.get(final_key, {}).get("content", "").strip()

        if not has_struct and not has_committed_projects and not final_content:
            continue

        # Section heading label
        label_map = {
            "work_experience": "Work Experience",
            "internships": "Internships",
            "technical_skills": "Technical Skills",
            "projects": "Projects",
            "education": "Education",
            "certifications": "Certifications",
            "awards": "Awards & Achievements",
            "interests": "Interests",
        }
        label = label_map.get(struct_key, struct_key.replace("_", " ").title())
        _add_heading(doc, label, rendering)

        # ── WORK EXPERIENCE / INTERNSHIPS ──────────────────────────────
        if struct_key in ("work_experience", "internships") and has_struct:
            for job in struct_data:
                if isinstance(job, dict) and job.get("job_title"):
                    _render_job_entry(doc, job, rendering)

        # ── EDUCATION ───────────────────────────────────────────────────
        elif struct_key == "education" and has_struct:
            for edu in struct_data:
                if isinstance(edu, dict):
                    _render_education_entry(doc, edu, rendering)

        # ── TECHNICAL SKILLS ────────────────────────────────────────────
        elif struct_key == "technical_skills" and has_struct:
            for cat in struct_data:
                if not isinstance(cat, dict):
                    continue
                category = cat.get("category", "Skills")
                items = cat.get("items", [])
                if not items:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(f"{category}: ")
                r.bold = True
                r.font.name = font_family
                r.font.size = Pt(font_size_body)
                vr = p.add_run(", ".join(items))
                vr.font.name = font_family
                vr.font.size = Pt(font_size_body)

        # ── PROJECTS ────────────────────────────────────────────────────
        elif struct_key == "projects":
            # Existing projects from resume
            if has_struct:
                for proj in struct_data:
                    if isinstance(proj, dict):
                        _render_project_entry(doc, proj, rendering)
            # Committed (suggested) projects from analysis
            for proj in (committed_projects or []):
                if isinstance(proj, dict):
                    _render_project_entry(doc, proj, rendering)

        # ── CERTIFICATIONS ──────────────────────────────────────────────
        elif struct_key == "certifications" and has_struct:
            for cert in struct_data:
                if isinstance(cert, dict) and cert.get("name"):
                    parts = [cert["name"]]
                    if cert.get("issuer"):
                        parts.append(cert["issuer"])
                    if cert.get("date"):
                        parts.append(cert["date"])
                    _add_bullet(doc, " | ".join(parts), rendering)

        # ── AWARDS ──────────────────────────────────────────────────────
        elif struct_key == "awards" and has_struct:
            for award in struct_data:
                if isinstance(award, dict) and award.get("title"):
                    _add_bullet(doc, award["title"], rendering)

        # ── INTERESTS ───────────────────────────────────────────────────
        elif struct_key == "interests" and has_struct:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            if isinstance(struct_data, list):
                vr = p.add_run(", ".join([i for i in struct_data if i]))
            else:
                vr = p.add_run(str(struct_data))
            vr.font.name = font_family
            vr.font.size = Pt(font_size_body)

        # ── FALLBACK: final_sections content ────────────────────────────
        elif final_content:
            for line in final_content.split("\n"):
                _add_bullet(doc, line, rendering)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def save_docx(docx_bytes: bytes, filename: str = "resume_optimized.docx") -> str:
    out_path = OUTPUT_DIR / filename
    with open(out_path, "wb") as f:
        f.write(docx_bytes)
    return str(out_path)


def convert_to_pdf(docx_path: str) -> str | None:
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        import subprocess
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), docx_path],
            capture_output=True, timeout=30
        )
        if result.returncode == 0:
            return pdf_path
    except Exception:
        pass
    return None
