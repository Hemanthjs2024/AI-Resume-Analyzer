import docx

def create_resume():
    doc = docx.Document()
    
    # Title
    doc.add_heading("John Doe", level=0)
    
    # Contact Info
    p = doc.add_paragraph()
    p.add_run("john.doe@example.com | +1 (123) 456-7890 | github.com/johndoe | linkedin.com/in/johndoe")
    
    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Highly motivated and results-oriented Software Engineer with a solid foundation in Python and backend web development. Experienced in building scalable APIs and deploying dockerized applications.")
    
    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, JavaScript\nFrontend: HTML, CSS\nBackend: Django, FastAPI\nDevOps & Tools: Git, GitHub, Docker")
    
    # Experience
    doc.add_heading("Experience", level=1)
    
    p_job1 = doc.add_paragraph()
    r1 = p_job1.add_run("Software Engineer")
    r1.bold = True
    p_job1.add_run(" | Acme Corp | 2021 – 2023")
    doc.add_paragraph("• Developed and maintained REST APIs using Python and Django, increasing system performance by 25%.")
    doc.add_paragraph("• Containerized internal tools using Docker, reducing onboarding setup time for new engineers.")
    doc.add_paragraph("• Collaborated with frontend developers to design clean API contracts and integrate user interfaces.")
    
    # Education
    doc.add_heading("Education", level=1)
    p_edu = doc.add_paragraph()
    r2 = p_edu.add_run("Bachelor of Science in Computer Science")
    r2.bold = True
    p_edu.add_run(" | Tech University | 2017 – 2021\nCGPA: 3.8/4.0")
    
    doc.save("test_resume.docx")
    print("test_resume.docx generated successfully!")

if __name__ == "__main__":
    create_resume()
